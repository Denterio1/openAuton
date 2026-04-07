"""
tools/file_parser.py
====================
Universal file parser – extracts text content from any file type.

Supported formats:
- Plain text: .txt, .csv, .json, .xml, .html, .md, .log, source code
- Documents: .pdf, .docx, .doc, .odt, .rtf
- Spreadsheets: .xlsx, .xls
- Images: .jpg, .png, .bmp, .tiff (uses OCR if needed, or captioning)
- Audio: .mp3, .wav, .flac, .m4a (transcription)
- Binary: .exe, .dll, .bin, .dat (extracts printable strings)
"""

from __future__ import annotations
import os
import re
import subprocess
import tempfile
from pathlib import Path
from typing import List, Optional, Tuple, Dict, Any
import logging

logger = logging.getLogger(__name__)

# Optional imports (graceful fallback)
try:
    import magic
    HAS_MAGIC = True
except ImportError:
    HAS_MAGIC = False

try:
    from pypdf import PdfReader
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    from PIL import Image
    import pytesseract
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

try:
    import whisper
    HAS_WHISPER = True
except ImportError:
    HAS_WHISPER = False


class UniversalFileParser:
    """Extracts text from any file by detecting its MIME type and using appropriate parser."""

    def __init__(self):
        self._init_magic()

    def _init_magic(self):
        if not HAS_MAGIC:
            logger.warning("python-magic not installed. Install with: pip install python-magic")
            self.magic = None
        else:
            self.magic = magic

    def extract_text(self, file_path: Path, max_chars: int = 100_000) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text content and metadata from any file.
        Returns (text, metadata).
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Detect MIME type
        mime_type = self._detect_mime(file_path)
        ext = file_path.suffix.lower()

        metadata = {
            "file_path": str(file_path),
            "file_size_bytes": file_path.stat().st_size,
            "mime_type": mime_type,
            "parser_used": "unknown",
        }

        # Route to appropriate parser
        if mime_type.startswith("text/") or ext in ['.txt', '.csv', '.json', '.xml', '.html', '.md', '.log', '.py', '.js', '.cpp', '.c', '.h', '.java', '.go', '.rs']:
            text, meta = self._parse_text(file_path)
            metadata["parser_used"] = "text"
        elif mime_type == "application/pdf" or ext == '.pdf':
            text, meta = self._parse_pdf(file_path)
            metadata["parser_used"] = "pdf"
        elif mime_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"] or ext in ['.docx', '.doc']:
            text, meta = self._parse_docx(file_path)
            metadata["parser_used"] = "docx"
        elif mime_type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/vnd.ms-excel"] or ext in ['.xlsx', '.xls']:
            text, meta = self._parse_excel(file_path)
            metadata["parser_used"] = "excel"
        elif mime_type.startswith("image/") or ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
            text, meta = self._parse_image(file_path)
            metadata["parser_used"] = "image"
        elif mime_type.startswith("audio/") or ext in ['.mp3', '.wav', '.flac', '.m4a']:
            text, meta = self._parse_audio(file_path)
            metadata["parser_used"] = "audio"
        else:
            # Fallback: binary – extract printable strings
            text, meta = self._parse_binary(file_path)
            metadata["parser_used"] = "binary"

        metadata.update(meta)
        # Truncate if too long
        if len(text) > max_chars:
            text = text[:max_chars] + f"\n... (truncated, original length {len(text)})"
        return text, metadata

    def _detect_mime(self, path: Path) -> str:
        if HAS_MAGIC:
            try:
                return self.magic.from_file(str(path), mime=True)
            except:
                pass
        # Fallback: guess by extension
        ext = path.suffix.lower()
        mime_map = {
            '.txt': 'text/plain', '.csv': 'text/csv', '.json': 'application/json',
            '.xml': 'application/xml', '.html': 'text/html', '.md': 'text/markdown',
            '.log': 'text/plain', '.py': 'text/x-python', '.js': 'text/javascript',
            '.pdf': 'application/pdf', '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword', '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.xls': 'application/vnd.ms-excel', '.jpg': 'image/jpeg', '.png': 'image/png',
            '.mp3': 'audio/mpeg', '.wav': 'audio/wav',
        }
        return mime_map.get(ext, 'application/octet-stream')

    def _parse_text(self, path: Path) -> Tuple[str, Dict]:
        try:
            with open(path, 'r', encoding='utf-8', errors='replace') as f:
                text = f.read()
            return text, {"encoding": "utf-8", "line_count": text.count('\n')}
        except UnicodeDecodeError:
            with open(path, 'r', encoding='latin-1', errors='replace') as f:
                text = f.read()
            return text, {"encoding": "latin-1", "line_count": text.count('\n')}

    def _parse_pdf(self, path: Path) -> Tuple[str, Dict]:
        if not HAS_PYPDF:
            raise ImportError("pypdf not installed. Install with: pip install pypdf")
        reader = PdfReader(path)
        text = "\n".join([page.extract_text() or "" for page in reader.pages])
        return text, {"num_pages": len(reader.pages)}

    def _parse_docx(self, path: Path) -> Tuple[str, Dict]:
        if not HAS_DOCX:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        doc = Document(path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text, {"num_paragraphs": len(doc.paragraphs)}

    def _parse_excel(self, path: Path) -> Tuple[str, Dict]:
        if not HAS_OPENPYXL:
            raise ImportError("openpyxl not installed. Install with: pip install openpyxl")
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        texts = []
        for sheet in wb.worksheets:
            sheet_text = []
            for row in sheet.iter_rows(values_only=True):
                row_text = " ".join([str(cell) for cell in row if cell is not None])
                if row_text.strip():
                    sheet_text.append(row_text)
            if sheet_text:
                texts.append(f"--- Sheet: {sheet.title} ---\n" + "\n".join(sheet_text))
        wb.close()
        return "\n\n".join(texts), {"num_sheets": len(wb.worksheets)}

    def _parse_image(self, path: Path) -> Tuple[str, Dict]:
        if not HAS_OCR:
            raise ImportError("OCR not available. Install: pip install pillow pytesseract")
        img = Image.open(path)
        text = pytesseract.image_to_string(img)
        return text, {"image_size": img.size, "mode": img.mode}

    def _parse_audio(self, path: Path) -> Tuple[str, Dict]:
        if not HAS_WHISPER:
            raise ImportError("whisper not installed. Install with: pip install openai-whisper")
        model = whisper.load_model("base")
        result = model.transcribe(str(path))
        text = result["text"]
        return text, {"language": result.get("language", "unknown"), "duration_seconds": result.get("segments", [{}])[-1].get("end", 0) if result.get("segments") else 0}

    def _parse_binary(self, path: Path) -> Tuple[str, Dict]:
        """Extract printable strings from binary file."""
        try:
            result = subprocess.run(['strings', str(path)], capture_output=True, text=True, timeout=10)
            strings = result.stdout
        except Exception:
            # Fallback: read raw and filter printable
            with open(path, 'rb') as f:
                data = f.read(1024*1024)  # first 1MB
            strings = ''.join(chr(b) for b in data if 32 <= b < 127 or b in (9,10,13))
        # Clean up
        lines = [line.strip() for line in strings.splitlines() if len(line.strip()) > 3]
        text = "\n".join(lines[:1000])  # limit
        return text, {"extracted_strings_count": len(lines)}


def get_file_text(file_path: Path, max_chars: int = 100_000) -> Tuple[str, Dict]:
    """Convenience function."""
    parser = UniversalFileParser()
    return parser.extract_text(file_path, max_chars)